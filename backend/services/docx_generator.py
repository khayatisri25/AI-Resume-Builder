import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.services.logging_service import logger

def add_p_border_bottom(paragraph, color_hex="4B5563", size_val="6"):
    """
    Applies a native bottom border (horizontal divider line) to a Word paragraph.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    
    # Check if pBdr already exists
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size_val) # size in 1/8 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """
    Sets inner margins (padding) for a table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_resume_docx(data: dict, output_path: str) -> str:
    """
    Generates an editable, high-quality, ATS-optimized Microsoft Word (.docx) resume.
    Returns the filepath of the generated document.
    """
    personal = data.get("personal_info", {})
    objective = data.get("objective", "")
    education = data.get("education", [])
    skills = data.get("skills", [])
    projects = data.get("projects", [])
    internships = data.get("internships", [])
    experience = data.get("work_experience", [])
    certifications = data.get("certifications", [])
    achievements = data.get("achievements", [])
    languages = data.get("languages", [])
    interests = data.get("interests", [])
    template = data.get("preferred_template", "modern").lower()

    logger.info(f"Generating DOCX at path: {output_path}")

    doc = Document()

    # Configure Margins (0.75 in for standard, 0.5 for minimal)
    margin_size = Inches(0.75)
    if template == "minimal":
        margin_size = Inches(0.5)
        
    for section in doc.sections:
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size

    # Base Colors
    primary_color_rgb = RGBColor(0, 0, 0) # Black
    color_hex = "000000"
    if template in ["modern", "two column"]:
        primary_color_rgb = RGBColor(15, 118, 110) # Teal (#0F766E)
        color_hex = "0F766E"
    elif template == "professional":
        primary_color_rgb = RGBColor(30, 58, 138) # Navy (#1E3A8A)
        color_hex = "1E3A8A"

    # Default Font styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial' if template not in ["classic", "harvard"] else 'Times New Roman'
    font.size = Pt(10)
    font.color.rgb = RGBColor(55, 65, 81) # Dark charcoal (#374151)

    # 1. Header (Name, Title, Contact Info)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template in ["classic", "harvard"] else WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_p.add_run(personal.get("full_name", ""))
    name_run.font.name = font.name
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = primary_color_rgb

    if personal.get("professional_title"):
        title_p = doc.add_paragraph()
        title_p.alignment = name_p.alignment
        # Adjust spacing
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(4)
        title_run = title_p.add_run(personal.get("professional_title"))
        title_run.font.name = font.name
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(107, 114, 128) # Gray (#6b7280)

    # Contact Details Paragraph
    contact_p = doc.add_paragraph()
    contact_p.alignment = name_p.alignment
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(12)
    
    parts = []
    if personal.get("phone"): parts.append(personal.get("phone"))
    if personal.get("email"): parts.append(personal.get("email"))
    if personal.get("address"): parts.append(personal.get("address"))
    if personal.get("linkedin"): parts.append(personal.get("linkedin"))
    if personal.get("github"): parts.append(personal.get("github"))
    if personal.get("portfolio"): parts.append(personal.get("portfolio"))

    separator = "  |  " if template in ["classic", "harvard"] else "  •  "
    contact_run = contact_p.add_run(separator.join(parts))
    contact_run.font.name = font.name
    contact_run.font.size = Pt(9.0)
    contact_run.font.color.rgb = RGBColor(75, 85, 99)

    def add_section_title(title_text):
        """
        Creates a section heading with a horizontal rule border.
        """
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(12)
        title_p.paragraph_format.space_after = Pt(4)
        title_p.paragraph_format.keep_with_next = True
        
        run = title_p.add_run(title_text.upper())
        run.font.name = font.name
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = primary_color_rgb
        
        add_p_border_bottom(title_p, color_hex=color_hex, size_val="6")

    # 2. Objective
    if objective:
        add_section_title("Professional Summary")
        obj_p = doc.add_paragraph(objective)
        obj_p.paragraph_format.space_after = Pt(8)
        obj_p.paragraph_format.line_spacing = 1.15

    # 3. Work Experience
    if experience:
        add_section_title("Work Experience")
        for exp in experience:
            # We align Company/Role on the left and Dates on the right using a borderless table
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(2.0)
            
            # Left cell (Role & Company)
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            r_role = p_left.add_run(f"{exp.get('role', '')} ")
            r_role.bold = True
            r_role.font.color.rgb = RGBColor(31, 41, 55)
            
            r_comp = p_left.add_run(f"| {exp.get('company', '')}")
            r_comp.bold = True
            
            # Right cell (Dates)
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            r_dates = p_right.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            r_dates.font.size = Pt(9.5)
            r_dates.font.color.rgb = RGBColor(75, 85, 99)
            
            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    set_cell_margins(cell, top=40, bottom=40, left=0, right=0)
            
            # Bullets
            desc = exp.get("description", "")
            bullets = desc.split('\n') if isinstance(desc, str) else desc
            for b in bullets:
                b_str = b.strip()
                if not b_str: continue
                if b_str.startswith("- ") or b_str.startswith("* "):
                    b_str = b_str[2:]
                elif b_str.startswith("-") or b_str.startswith("*"):
                    b_str = b_str[1:]
                    
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.paragraph_format.space_before = Pt(0)
                bullet_p.paragraph_format.space_after = Pt(2)
                run = bullet_p.add_run(b_str)
                run.font.size = Pt(9.5)

            # Spacing between roles
            spacer_p = doc.add_paragraph()
            spacer_p.paragraph_format.space_before = Pt(0)
            spacer_p.paragraph_format.space_after = Pt(2)

    # 4. Internships
    if internships:
        add_section_title("Internships")
        for intern in internships:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(2.0)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            r_role = p_left.add_run(f"{intern.get('role', '')} ")
            r_role.bold = True
            r_role.font.color.rgb = RGBColor(31, 41, 55)
            
            r_comp = p_left.add_run(f"| {intern.get('company', '')}")
            r_comp.bold = True
            
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            r_dates = p_right.add_run(f"{intern.get('start_date', '')} - {intern.get('end_date', '')}")
            r_dates.font.size = Pt(9.5)
            r_dates.font.color.rgb = RGBColor(75, 85, 99)
            
            for row in table.rows:
                for cell in row.cells:
                    set_cell_margins(cell, top=40, bottom=40, left=0, right=0)
            
            desc = intern.get("description", "")
            bullets = desc.split('\n') if isinstance(desc, str) else desc
            for b in bullets:
                b_str = b.strip()
                if not b_str: continue
                if b_str.startswith("- ") or b_str.startswith("* "):
                    b_str = b_str[2:]
                elif b_str.startswith("-") or b_str.startswith("*"):
                    b_str = b_str[1:]
                    
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.paragraph_format.space_before = Pt(0)
                bullet_p.paragraph_format.space_after = Pt(2)
                run = bullet_p.add_run(b_str)
                run.font.size = Pt(9.5)

            spacer_p = doc.add_paragraph()
            spacer_p.paragraph_format.space_before = Pt(0)
            spacer_p.paragraph_format.space_after = Pt(2)

    # 5. Projects
    if projects:
        add_section_title("Projects")
        for proj in projects:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.2)
            table.columns[1].width = Inches(1.8)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            
            r_name = p_left.add_run(proj.get("project_name", ""))
            r_name.bold = True
            r_name.font.color.rgb = RGBColor(31, 41, 55)
            
            if proj.get("technologies"):
                r_tech = p_left.add_run(f" ({proj.get('technologies')})")
                r_tech.font.size = Pt(9.0)
                r_tech.font.color.rgb = RGBColor(107, 114, 128)
                
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            
            links = []
            if proj.get("github_link"): links.append("GitHub")
            if proj.get("live_link"): links.append("Live")
            r_links = p_right.add_run(" | ".join(links))
            r_links.font.size = Pt(9.0)
            r_links.font.color.rgb = RGBColor(59, 130, 246)
            
            for row in table.rows:
                for cell in row.cells:
                    set_cell_margins(cell, top=40, bottom=40, left=0, right=0)
            
            desc = proj.get("description", "")
            bullets = desc.split('\n') if isinstance(desc, str) else desc
            for b in bullets:
                b_str = b.strip()
                if not b_str: continue
                if b_str.startswith("- ") or b_str.startswith("* "):
                    b_str = b_str[2:]
                elif b_str.startswith("-") or b_str.startswith("*"):
                    b_str = b_str[1:]
                    
                bullet_p = doc.add_paragraph(style='List Bullet')
                bullet_p.paragraph_format.space_before = Pt(0)
                bullet_p.paragraph_format.space_after = Pt(2)
                run = bullet_p.add_run(b_str)
                run.font.size = Pt(9.5)

            spacer_p = doc.add_paragraph()
            spacer_p.paragraph_format.space_before = Pt(0)
            spacer_p.paragraph_format.space_after = Pt(2)

    # 6. Education
    if education:
        add_section_title("Education")
        for edu in education:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(2.0)
            
            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            
            inst_parts = []
            if edu.get("college"): inst_parts.append(edu.get("college"))
            if edu.get("university"): inst_parts.append(edu.get("university"))
            inst_str = ", ".join(inst_parts)
            
            r_deg = p_left.add_run(f"{edu.get('degree', '')} ")
            r_deg.bold = True
            r_deg.font.color.rgb = RGBColor(31, 41, 55)
            
            r_inst = p_left.add_run(f"| {inst_str}")
            
            if edu.get("cgpa"):
                p_left.add_run(f" (CGPA: {edu.get('cgpa')})")
                
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(2)
            r_dates = p_right.add_run(f"{edu.get('start_year', '')} - {edu.get('end_year', '')}")
            r_dates.font.size = Pt(9.5)
            r_dates.font.color.rgb = RGBColor(75, 85, 99)
            
            for row in table.rows:
                for cell in row.cells:
                    set_cell_margins(cell, top=40, bottom=40, left=0, right=0)

    # 7. Skills
    if skills:
        add_section_title("Skills")
        skills_p = doc.add_paragraph()
        skills_p.paragraph_format.space_after = Pt(6)
        skills_run = skills_p.add_run(", ".join(skills))
        skills_run.font.size = Pt(9.5)

    # 8. Certifications & Achievements
    if certifications or achievements:
        if certifications and achievements:
            add_section_title("Certifications & Achievements")
        elif certifications:
            add_section_title("Certifications")
        else:
            add_section_title("Achievements")
            
        for cert in certifications:
            cert_p = doc.add_paragraph(style='List Bullet')
            cert_p.paragraph_format.space_after = Pt(2)
            run = cert_p.add_run(f"{cert.get('name', '')} ({cert.get('issuer', '')}) - {cert.get('year', '')}")
            run.font.size = Pt(9.5)
            
        for ach in achievements:
            ach_p = doc.add_paragraph(style='List Bullet')
            ach_p.paragraph_format.space_after = Pt(2)
            run = ach_p.add_run(f"{ach.get('title', '')}: {ach.get('description', '')}")
            run.font.size = Pt(9.5)

    # 9. Languages & Interests
    if languages or interests:
        if languages and interests:
            add_section_title("Languages & Interests")
        elif languages:
            add_section_title("Languages")
        else:
            add_section_title("Interests")
            
        lang_str_list = []
        for lang in languages:
            name = lang.get("language", "")
            prof = lang.get("proficiency", "")
            lang_str_list.append(f"{name} ({prof})" if prof else name)
            
        final_parts = []
        if lang_str_list:
            final_parts.append(f"Languages: {', '.join(lang_str_list)}")
        if interests:
            final_parts.append(f"Interests: {', '.join(interests)}")
            
        lang_p = doc.add_paragraph()
        lang_p.paragraph_format.space_after = Pt(6)
        lang_run = lang_p.add_run("  |  ".join(final_parts))
        lang_run.font.size = Pt(9.5)

    doc.save(output_path)
    logger.info(f"Successfully generated DOCX file at {output_path}")
    return output_path
